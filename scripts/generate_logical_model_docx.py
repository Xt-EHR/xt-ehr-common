#!/usr/bin/env python3
"""Generate a Word document with tables from StructureDefinition JSON files."""

# Usage:
#   python3 scripts/generate_logical_model_docx.py
#   python3 scripts/generate_logical_model_docx.py fsh-generated/resources/StructureDefinition-EHDSAddress.json
#   python3 scripts/generate_logical_model_docx.py --input-path fsh-generated/resources -o out/logical-models.docx
#
# Input parameters:
#   input
#     Optional positional path to a single StructureDefinition `.json` file or a
#     directory containing StructureDefinition JSON files. Defaults to
#     `fsh-generated/resources`.
#   -i, --input-path
#     Optional named path to a single StructureDefinition `.json` file or a
#     directory containing StructureDefinition JSON files. If provided, it
#     overrides the positional `input` argument.
#   -o, --output
#     Optional output path for the generated `.docx` file.
#
# Notes:
#   The script resolves the `baseDefinition` chain using local JSON files from
#   the same directory, so inherited elements are included in the output table.

from __future__ import annotations

import argparse
import json
import heapq
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


HEADER_FILL = "D9D9D9"
PAGE_FIT_WIDTHS = [0.7, 1.6, 3.0, 1.7, 1.0, 1.6]


@dataclass
class ElementRow:
    level: int
    name: str
    description: str = ""
    data_types: list[str] = field(default_factory=list)
    cardinality: str = ""
    binding_description: str = ""

    @property
    def binding(self) -> str:
        return self.binding_description

    @property
    def data_type(self) -> str:
        return "\n".join(self.data_types)


@dataclass
class LogicalModel:
    logical_name: str
    elements: list[ElementRow] = field(default_factory=list)


@dataclass
class StructureDefinitionRepository:
    by_url: dict[str, dict]
    by_id: dict[str, dict]
    by_name: dict[str, dict]

    def get(self, reference: str) -> dict | None:
        if not reference:
            return None
        return self.by_url.get(reference) or self.by_id.get(reference) or self.by_name.get(reference)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate a Word document with tables from StructureDefinition JSON files."
    )
    parser.add_argument(
        "input",
        nargs="?",
        default="fsh-generated/resources",
        help="StructureDefinition JSON file or directory. Defaults to fsh-generated/resources.",
    )
    parser.add_argument(
        "-i",
        "--input-path",
        dest="input_path",
        help="StructureDefinition JSON file or directory. Overrides the positional input argument.",
    )
    parser.add_argument(
        "-o",
        "--output",
        default="out/logical-models.docx",
        help="Output .docx path. Defaults to out/logical-models.docx.",
    )
    return parser.parse_args()


def iter_input_files(input_path: Path) -> Iterable[Path]:
    if not input_path.exists():
        raise SystemExit(f"Input path does not exist: {input_path}")
    if input_path.is_file():
        if input_path.suffix.lower() != ".json":
            raise SystemExit(f"Input file must have a .json extension: {input_path}")
        if "Obligations" in input_path.name:
            raise SystemExit(f"Input file is excluded because it contains 'Obligations': {input_path}")
        yield input_path
        return
    if not input_path.is_dir():
        raise SystemExit(f"Input path must be a directory or .json file: {input_path}")
    for path in sorted(input_path.glob("StructureDefinition-*.json")):
        if path.is_file() and "Obligations" not in path.name:
            yield path


def build_repository(search_dir: Path) -> StructureDefinitionRepository:
    by_url: dict[str, dict] = {}
    by_id: dict[str, dict] = {}
    by_name: dict[str, dict] = {}

    for path in sorted(search_dir.glob("StructureDefinition-*.json")):
        if not path.is_file():
            continue
        obj = json.loads(path.read_text(encoding="utf-8"))
        if obj.get("resourceType") != "StructureDefinition":
            continue
        if obj.get("url"):
            by_url[obj["url"]] = obj
        if obj.get("id"):
            by_id[obj["id"]] = obj
        if obj.get("name"):
            by_name[obj["name"]] = obj

    return StructureDefinitionRepository(by_url=by_url, by_id=by_id, by_name=by_name)


def merge_elements(base_elements: list[dict], differential_elements: list[dict]) -> list[dict]:
    merged_by_key: dict[str, dict] = {}
    base_order: list[str] = []
    diff_order: list[str] = []

    for element in base_elements:
        key = element_merge_key(element)
        if key is None:
            continue
        merged_by_key[key] = deepcopy(element)
        base_order.append(key)

    for element in differential_elements:
        key = element_merge_key(element)
        if key is None:
            continue
        if key in merged_by_key:
            merged_by_key[key].update(deepcopy(element))
        else:
            merged_by_key[key] = deepcopy(element)
        diff_order.append(key)

    ordered_keys = stable_topological_order(base_order, diff_order)
    return [merged_by_key[key] for key in ordered_keys if key in merged_by_key]


def stable_topological_order(base_order: list[str], diff_order: list[str]) -> list[str]:
    all_keys = list(dict.fromkeys(base_order + diff_order))
    graph: dict[str, set[str]] = {key: set() for key in all_keys}
    indegree: dict[str, int] = {key: 0 for key in all_keys}
    overall_rank = {key: index for index, key in enumerate(all_keys)}

    def add_edge(first: str, second: str) -> None:
        if first == second:
            return
        if second not in graph[first]:
            graph[first].add(second)
            indegree[second] += 1

    def add_edges(order: list[str]) -> None:
        for first, second in zip(order, order[1:]):
            add_edge(first, second)

    add_edges(base_order)
    add_edges(diff_order)

    base_only = set(base_order)
    for key in diff_order:
        if key in base_only:
            continue
        boundary = find_parent_subtree_boundary(base_order, key)
        if boundary:
            add_edge(key, boundary)

    base_rank = {key: index for index, key in enumerate(base_order)}
    diff_rank = {key: index for index, key in enumerate(diff_order)}

    heap: list[tuple[tuple[int, int, int], str]] = []
    for position, key in enumerate(all_keys):
        if indegree[key] == 0:
            priority = (base_rank.get(key, 10**9), diff_rank.get(key, 10**9), position)
            heapq.heappush(heap, (priority, key))

    ordered: list[str] = []
    while heap:
        _, key = heapq.heappop(heap)
        ordered.append(key)
        for neighbor in graph[key]:
            indegree[neighbor] -= 1
            if indegree[neighbor] == 0:
                priority = (base_rank.get(neighbor, 10**9), diff_rank.get(neighbor, 10**9), overall_rank[neighbor])
                heapq.heappush(heap, (priority, neighbor))

    if len(ordered) != len(all_keys):
        raise SystemExit("Could not resolve StructureDefinition element order.")

    return ordered


def parent_key(key: str) -> str:
    if "." not in key:
        return ""
    return key.rsplit(".", 1)[0]


def find_parent_subtree_boundary(base_order: list[str], key: str) -> str | None:
    parent = parent_key(key)
    if not parent:
        return None
    if parent not in base_order:
        return find_parent_subtree_boundary(base_order, parent)

    parent_index = base_order.index(parent)
    prefix = f"{parent}."
    index = parent_index + 1
    while index < len(base_order):
        candidate = base_order[index]
        if candidate.startswith(prefix):
            index += 1
            continue
        return candidate
    return None


def element_merge_key(element: dict) -> str | None:
    path_value = element.get("path") or element.get("id")
    if not path_value:
        return None
    segments = path_value.split(".")
    return ".".join(segments[1:])


def resolve_elements(definition: dict, repository: StructureDefinitionRepository, seen: set[str] | None = None) -> list[dict]:
    seen = seen or set()
    key = definition.get("url") or definition.get("id") or definition.get("name")
    if key:
        if key in seen:
            raise SystemExit(f"Cyclic baseDefinition chain detected at {key}")
        seen = set(seen)
        seen.add(key)

    base_elements: list[dict] = []
    base_reference = definition.get("baseDefinition")
    if base_reference:
        base_definition = repository.get(base_reference)
        if base_definition is not None:
            base_elements = resolve_elements(base_definition, repository, seen)

    own_elements = definition.get("differential", {}).get("element", []) or []
    return merge_elements(base_elements, own_elements)


def sanitize_text(value: str) -> str:
    return value.replace("\\", "")


def format_type_code(code: str) -> str:
    if not code:
        return ""
    return sanitize_text(code.rstrip("/").split("/")[-1])


def format_types(element: dict) -> str:
    codes = [format_type_code(item.get("code", "")) for item in element.get("type", [])]
    return [code for code in codes if code]


def format_cardinality(element: dict) -> str:
    minimum = element.get("min")
    maximum = element.get("max")
    if minimum is None and maximum is None:
        return ""
    return f"{minimum if minimum is not None else ''}..{maximum if maximum is not None else ''}"


def format_description(element: dict) -> str:
    definition = sanitize_text((element.get("definition") or "").strip())
    short = sanitize_text((element.get("short") or "").strip())
    return definition or short


def path_level_and_name(path_value: str) -> tuple[int, str]:
    segments = path_value.split(".")
    return len(segments) - 1, sanitize_text(segments[-1])


def parse_structure_definition(path: Path, repository: StructureDefinitionRepository) -> LogicalModel | None:
    definition = json.loads(path.read_text(encoding="utf-8"))
    if definition.get("resourceType") != "StructureDefinition":
        return None
    if definition.get("kind") != "logical":
        return None

    logical_name = definition.get("id") or definition.get("name") or path.stem
    resolved_elements = resolve_elements(definition, repository)
    rows: list[ElementRow] = []

    for element in resolved_elements:
        path_value = element.get("path", "")
        if not path_value or "." not in path_value:
            continue
        level, name = path_level_and_name(path_value)
        rows.append(
            ElementRow(
                level=level,
                name=name,
                description=format_description(element),
                data_types=format_types(element),
                cardinality=format_cardinality(element),
                binding_description=sanitize_text((element.get("binding", {}) or {}).get("description", "")),
            )
        )

    return LogicalModel(logical_name=logical_name, elements=rows)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_width(cell, inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_top_left(cell) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def get_page_fit_widths(document: Document) -> list[float]:
    section = document.sections[-1]
    usable_width = (section.page_width - section.left_margin - section.right_margin) / 914400
    base_total = sum(PAGE_FIT_WIDTHS)
    scale = usable_width / base_total
    return [width * scale for width in PAGE_FIT_WIDTHS]


def add_model_table(document: Document, model: LogicalModel) -> None:
    document.add_heading(model.logical_name, level=1)

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = get_page_fit_widths(document)
    headers = [
        "Level",
        "Element name",
        "Element description",
        "Data type",
        "Cardinality",
        "Binding",
    ]

    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, heading in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_width(cell, widths[index])
        shade_cell(cell, HEADER_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = paragraph.add_run(heading)
        run.bold = True
        set_cell_top_left(cell)

    for element in model.elements:
        row = table.add_row()
        values = [
            "+" * element.level,
            element.name,
            element.description,
            element.data_type,
            element.cardinality,
            element.binding,
        ]
        for index, value in enumerate(values):
            cell = row.cells[index]
            set_cell_width(cell, widths[index])
            cell.text = value
            set_cell_top_left(cell)


def build_document(models: list[LogicalModel], output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos Narrow"
    normal_style.font.size = Pt(9)

    for index, model in enumerate(models):
        if index > 0:
            document.add_section(WD_SECTION.NEW_PAGE)
        add_model_table(document, model)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    args = parse_args()
    input_path = Path(args.input_path or args.input).expanduser()
    output_path = Path(args.output)

    search_dir = input_path.parent if input_path.is_file() else input_path
    repository = build_repository(search_dir)

    models: list[LogicalModel] = []
    for file_path in iter_input_files(input_path):
        model = parse_structure_definition(file_path, repository)
        if model is not None:
            models.append(model)

    if not models:
        raise SystemExit(f"No logical StructureDefinition files found in {input_path}")

    build_document(models, output_path)
    print(f"Created {output_path} from {len(models)} StructureDefinition file(s).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
