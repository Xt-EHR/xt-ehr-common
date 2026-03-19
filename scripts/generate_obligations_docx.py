#!/usr/bin/env python3
"""Generate a Word document with tables from obligation StructureDefinition JSON files."""

# Usage:
#   python3 scripts/generate_obligations_docx.py
#   python3 scripts/generate_obligations_docx.py fsh-generated/resources/StructureDefinition-EHDSLaboratoryReportObligations.json
#   python3 scripts/generate_obligations_docx.py --input-path fsh-generated/resources -o out/obligations.docx
#
# Input parameters:
#   input
#     Optional positional path to a single obligation StructureDefinition `.json`
#     file or a directory containing them. Defaults to `fsh-generated/resources`.
#   -i, --input-path
#     Optional named path to a single obligation StructureDefinition `.json` file
#     or a directory containing them. If provided, it overrides the positional
#     `input` argument.
#   -o, --output
#     Optional output path for the generated `.docx` file.

from __future__ import annotations

import argparse
import json
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


HEADER_FILL = "D9D9D9"
PAGE_FIT_WIDTHS = [0.6, 1.5, 0.8, 1.2, 2.2, 1.2, 2.2]
OBLIGATION_URL = "http://hl7.org/fhir/StructureDefinition/obligation"
PRODUCER_ACTOR = "https://www.xt-ehr.eu/specifications/fhir/actor-producer"
CONSUMER_ACTOR = "https://www.xt-ehr.eu/specifications/fhir/actor-consumer"


@dataclass
class ObligationEntry:
    code: str = ""
    documentation: str = ""


@dataclass
class ObligationRow:
    level: int
    name: str
    cardinality: str = ""
    producer: list[ObligationEntry] = field(default_factory=list)
    consumer: list[ObligationEntry] = field(default_factory=list)


@dataclass
class LogicalModel:
    logical_name: str
    elements: list[ObligationRow] = field(default_factory=list)


@dataclass
class StructureDefinitionRepository:
    by_url: dict[str, dict]
    by_id: dict[str, dict]
    by_name: dict[str, dict]

    def get(self, reference: str) -> dict | None:
        if not reference:
            return None
        return self.by_url.get(reference) or self.by_id.get(reference) or self.by_name.get(reference)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate a Word document with tables from obligation StructureDefinition JSON files."
    )
    parser.add_argument(
        "input",
        nargs="?",
        default="fsh-generated/resources",
        help="Obligation StructureDefinition JSON file or directory. Defaults to fsh-generated/resources.",
    )
    parser.add_argument(
        "-i",
        "--input-path",
        dest="input_path",
        help="Obligation StructureDefinition JSON file or directory. Overrides the positional input argument.",
    )
    parser.add_argument(
        "-o",
        "--output",
        default="out/obligations.docx",
        help="Output .docx path. Defaults to out/obligations.docx.",
    )
    return parser.parse_args()


def iter_input_files(input_path: Path) -> Iterable[Path]:
    if not input_path.exists():
        raise SystemExit(f"Input path does not exist: {input_path}")
    if input_path.is_file():
        if input_path.suffix.lower() != ".json":
            raise SystemExit(f"Input file must have a .json extension: {input_path}")
        if "Obligations" not in input_path.name:
            raise SystemExit(f"Input file must contain 'Obligations' in its name: {input_path}")
        yield input_path
        return
    if not input_path.is_dir():
        raise SystemExit(f"Input path must be a directory or .json file: {input_path}")
    for path in sorted(input_path.glob("StructureDefinition-*Obligations*.json")):
        if path.is_file():
            yield path


def build_repository(search_dir: Path) -> StructureDefinitionRepository:
    by_url: dict[str, dict] = {}
    by_id: dict[str, dict] = {}
    by_name: dict[str, dict] = {}

    for path in sorted(search_dir.glob("StructureDefinition-*.json")):
        if not path.is_file():
            continue
        obj = json.loads(path.read_text(encoding="utf-8"))
        if obj.get("resourceType") != "StructureDefinition":
            continue
        if obj.get("url"):
            by_url[obj["url"]] = obj
        if obj.get("id"):
            by_id[obj["id"]] = obj
        if obj.get("name"):
            by_name[obj["name"]] = obj

    return StructureDefinitionRepository(by_url=by_url, by_id=by_id, by_name=by_name)


def element_merge_key(element: dict) -> str | None:
    path_value = element.get("path") or element.get("id")
    if not path_value:
        return None
    head = path_value.split(":", 1)[0]
    segments = head.split(".")
    return ".".join(segments[1:])


def merge_elements(base_elements: list[dict], differential_elements: list[dict]) -> list[dict]:
    merged: list[dict] = [deepcopy(element) for element in base_elements]
    index_by_key = {
        element_merge_key(element): index for index, element in enumerate(merged) if element_merge_key(element) is not None
    }

    for diff in differential_elements:
        key = element_merge_key(diff)
        if key is not None and key in index_by_key:
            merged[index_by_key[key]].update(deepcopy(diff))
            continue
        merged.append(deepcopy(diff))
        if key is not None:
            index_by_key[key] = len(merged) - 1

    return merged


def resolve_elements(definition: dict, repository: StructureDefinitionRepository, seen: set[str] | None = None) -> list[dict]:
    seen = seen or set()
    key = definition.get("url") or definition.get("id") or definition.get("name")
    if key:
        if key in seen:
            raise SystemExit(f"Cyclic baseDefinition chain detected at {key}")
        seen = set(seen)
        seen.add(key)

    base_elements: list[dict] = []
    base_reference = definition.get("baseDefinition")
    if base_reference:
        base_definition = repository.get(base_reference)
        if base_definition is not None:
            base_elements = resolve_elements(base_definition, repository, seen)

    own_elements = definition.get("differential", {}).get("element", []) or []
    return merge_elements(base_elements, own_elements)


def format_cardinality(element: dict) -> str:
    minimum = element.get("min")
    maximum = element.get("max")
    if minimum is None and maximum is None:
        return ""
    return f"{minimum if minimum is not None else ''}..{maximum if maximum is not None else ''}"


def sanitize_text(value: str) -> str:
    return value.replace("\\", "")


def element_level_and_name(element: dict) -> tuple[int, str]:
    id_value = element.get("id") or element.get("path") or ""
    base_part = id_value.split(":", 1)[0]
    segments = base_part.split(".")
    level = len(segments) - 1
    name = segments[-1]
    if ":" in id_value:
        name = id_value.split(".")[-1]
    return level, sanitize_text(name)


def parse_obligation_extension(extension: dict) -> tuple[str, ObligationEntry] | None:
    if extension.get("url") != OBLIGATION_URL:
        return None

    actor = ""
    code = ""
    documentation = ""
    for child in extension.get("extension", []):
        url = child.get("url")
        if url == "actor":
            actor = child.get("valueCanonical", "")
        elif url == "code":
            code = sanitize_text(child.get("valueCode", ""))
        elif url == "documentation":
            documentation = sanitize_text(child.get("valueMarkdown", ""))

    if actor == PRODUCER_ACTOR:
        return "producer", ObligationEntry(code=code, documentation=documentation)
    if actor == CONSUMER_ACTOR:
        return "consumer", ObligationEntry(code=code, documentation=documentation)
    return None


def parse_obligation_model(path: Path, repository: StructureDefinitionRepository) -> LogicalModel | None:
    definition = json.loads(path.read_text(encoding="utf-8"))
    if definition.get("resourceType") != "StructureDefinition":
        return None
    if "Obligations" not in path.name:
        return None

    base_definition = repository.get(definition.get("baseDefinition", ""))
    resolved_base_elements = resolve_elements(base_definition, repository) if base_definition is not None else []
    base_by_key = {element_merge_key(element): element for element in resolved_base_elements if element_merge_key(element) is not None}
    resolved_elements = resolve_elements(definition, repository)
    resolved_by_key = {element_merge_key(element): element for element in resolved_elements if element_merge_key(element) is not None}

    rows_by_key: dict[str, ObligationRow] = {}
    included_keys: set[str] = set()
    for element in definition.get("differential", {}).get("element", []):
        parsed_extensions = [parse_obligation_extension(ext) for ext in element.get("extension", [])]
        parsed_extensions = [item for item in parsed_extensions if item is not None]
        if not parsed_extensions:
            continue

        key = element_merge_key(element)
        if key is None:
            continue

        base_element = base_by_key.get(key)
        level, name = element_level_and_name(element)
        row = rows_by_key.setdefault(
            key,
            ObligationRow(
                level=level,
                name=name,
                cardinality=format_cardinality(base_element or element),
            ),
        )
        included_keys.add(key)

        for actor, obligation in parsed_extensions:
            if actor == "producer":
                row.producer.append(obligation)
            elif actor == "consumer":
                row.consumer.append(obligation)

        parent_key = key.rsplit(".", 1)[0] if "." in key else ""
        while parent_key:
            if parent_key in resolved_by_key:
                included_keys.add(parent_key)
            parent_key = parent_key.rsplit(".", 1)[0] if "." in parent_key else ""

    rows: list[ObligationRow] = []
    for element in resolved_elements:
        key = element_merge_key(element)
        if key is None or key not in included_keys:
            continue
        existing_row = rows_by_key.get(key)
        if existing_row is not None:
            rows.append(existing_row)
            continue

        level, name = element_level_and_name(element)
        rows.append(
            ObligationRow(
                level=level,
                name=name,
                cardinality=format_cardinality(element),
            )
        )

    logical_name = definition.get("id") or definition.get("name") or path.stem
    return LogicalModel(logical_name=logical_name, elements=rows)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_width(cell, inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_vertical_merge(cell, restart: bool) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    v_merge = tc_pr.find(qn("w:vMerge"))
    if v_merge is None:
        v_merge = OxmlElement("w:vMerge")
        tc_pr.append(v_merge)
    v_merge.set(qn("w:val"), "restart" if restart else "continue")


def set_cell_top_left(cell) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def get_page_fit_widths(document: Document) -> list[float]:
    section = document.sections[-1]
    usable_width = (section.page_width - section.left_margin - section.right_margin) / 914400
    base_total = sum(PAGE_FIT_WIDTHS)
    scale = usable_width / base_total
    return [width * scale for width in PAGE_FIT_WIDTHS]


def add_row(table, widths: list[float], values: list[str]) -> None:
    row = table.add_row()
    for index, value in enumerate(values):
        cell = row.cells[index]
        set_cell_width(cell, widths[index])
        cell.text = value
        set_cell_top_left(cell)


def add_obligation_rows(table, widths: list[float], element: ObligationRow) -> None:
    total_rows = max(len(element.producer), len(element.consumer), 1)
    for index in range(total_rows):
        producer = element.producer[index] if index < len(element.producer) else ObligationEntry()
        consumer = element.consumer[index] if index < len(element.consumer) else ObligationEntry()
        add_row(
            table,
            widths,
            [
                "+" * element.level if index == 0 else "",
                element.name if index == 0 else "",
                element.cardinality if index == 0 else "",
                producer.code,
                producer.documentation,
                consumer.code,
                consumer.documentation,
            ],
        )
        row = table.rows[-1]
        if total_rows > 1:
            for column_index in range(3):
                set_vertical_merge(row.cells[column_index], restart=(index == 0))


def add_model_table(document: Document, model: LogicalModel) -> None:
    document.add_heading(model.logical_name, level=1)

    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = get_page_fit_widths(document)
    headers = [
        "Level",
        "Element name",
        "Cardinality",
        "Producer",
        "Producer documentation",
        "Consumer",
        "Consumer documentation",
    ]

    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, heading in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_width(cell, widths[index])
        shade_cell(cell, HEADER_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = paragraph.add_run(heading)
        run.bold = True
        set_cell_top_left(cell)

    for element in model.elements:
        add_obligation_rows(table, widths, element)


def build_document(models: list[LogicalModel], output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Aptos Narrow"
    normal_style.font.size = Pt(9)

    for index, model in enumerate(models):
        if index > 0:
            document.add_section(WD_SECTION.NEW_PAGE)
        add_model_table(document, model)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    args = parse_args()
    input_path = Path(args.input_path or args.input).expanduser()
    output_path = Path(args.output)

    search_dir = input_path.parent if input_path.is_file() else input_path
    repository = build_repository(search_dir)

    models: list[LogicalModel] = []
    for file_path in iter_input_files(input_path):
        model = parse_obligation_model(file_path, repository)
        if model is not None:
            models.append(model)

    if not models:
        raise SystemExit(f"No obligation StructureDefinition files found in {input_path}")

    build_document(models, output_path)
    print(f"Created {output_path} from {len(models)} obligation StructureDefinition file(s).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
